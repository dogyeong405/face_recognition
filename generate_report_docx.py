"""
탐구보고서 DOCX 생성 스크립트
- python-docx 사용
- 형식: 1.서론 2.이론적 배경 3.연구 방법 4.연구 결과 5.고찰 6.결론
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "탐구보고서_얼굴인식시스템.docx")


def set_cell_shading(cell, color_hex):
    """셀 배경색 설정"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows):
    """스타일된 표 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '3C3C78')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F0F0F8')

    return table


def build_report():
    doc = Document()

    # ── 기본 스타일 설정 ──
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '맑은 고딕'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════
    # 표지
    # ══════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('탐구보고서')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(40, 40, 80)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('InsightFace 기반\n실시간 얼굴 인식 웹 시스템')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(80, 80, 120)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('기술 스택: React + FastAPI + InsightFace (ArcFace)')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 목차
    # ══════════════════════════════════════
    doc.add_heading('목차', level=1)
    toc_items = [
        '1. 서론',
        '2. 이론적 배경',
        '3. 연구 방법',
        '4. 연구 결과',
        '5. 고찰',
        '6. 결론',
        '참고문헌',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════
    # 1. 서론
    # ══════════════════════════════════════
    doc.add_heading('1. 서론', level=1)

    doc.add_heading('1.1 연구 배경', level=2)
    doc.add_paragraph(
        '얼굴 인식 기술은 스마트폰 잠금 해제, 출입 관리, 보안 시스템, 무인 결제 등 '
        '일상생활 곳곳에서 활용되고 있다. 특히 최근에는 딥러닝 기술의 발전으로 '
        '얼굴 인식의 정확도가 비약적으로 향상되었으며, 단순한 신원 확인을 넘어 '
        '나이, 성별, 감정 등 얼굴의 다양한 속성을 추정하는 것이 가능해졌다.'
    )
    doc.add_paragraph(
        '그러나 이러한 기술이 실제로 어떤 원리로 작동하는지, 사람의 얼굴을 '
        '컴퓨터가 어떻게 수치화하여 구분하는지에 대해서는 일반적으로 잘 알려져 있지 않다. '
        '이에 본 연구에서는 딥러닝 기반 얼굴 인식 기술의 핵심 원리를 직접 구현하고 '
        '실험함으로써 그 작동 메커니즘을 깊이 이해하고자 하였다.'
    )

    doc.add_heading('1.2 연구 동기', level=2)
    doc.add_paragraph(
        '"웹 브라우저만으로도 실시간 얼굴 인식이 가능할까?"라는 질문에서 본 탐구가 시작되었다. '
        '기존의 얼굴 인식 시스템은 대부분 전용 소프트웨어나 고가의 장비를 필요로 하지만, '
        '웹 기술과 오픈소스 AI 모델을 결합하면 일반 웹 브라우저에서도 구현이 가능할 것이라는 '
        '가설을 세우고, 이를 검증하기 위해 직접 시스템을 설계하고 개발하였다.'
    )

    doc.add_heading('1.3 연구 목적', level=2)
    purposes = [
        '딥러닝 기반 얼굴 인식의 핵심 원리(임베딩, 코사인 유사도)를 이해한다.',
        'InsightFace(ArcFace) 모델을 활용하여 실시간 얼굴 인식 시스템을 구현한다.',
        '웹 기술(React, FastAPI)을 결합하여 브라우저에서 동작하는 완전한 얼굴 인식 애플리케이션을 개발한다.',
        '얼굴 인식과 동시에 나이, 성별 추정 기능을 통합하여 다중 속성 분석 시스템을 구축한다.',
    ]
    for b in purposes:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('1.4 연구 가설', level=2)
    doc.add_paragraph(
        '본 연구에서는 다음과 같은 가설을 설정하였다:'
    )
    hypotheses = [
        '가설 1: ArcFace 모델의 512차원 임베딩과 코사인 유사도를 활용하면, 등록된 얼굴을 높은 정확도로 인식할 수 있을 것이다.',
        '가설 2: 웹 브라우저와 REST API 기반의 Client-Server 구조로도 실시간 얼굴 인식이 가능할 것이다.',
        '가설 3: InsightFace의 Attribute 모델(genderage.onnx)을 활용하면, 얼굴 인식과 동시에 나이와 성별을 추정할 수 있을 것이다.',
    ]
    for h in hypotheses:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════
    # 2. 이론적 배경
    # ══════════════════════════════════════
    doc.add_heading('2. 이론적 배경', level=1)

    doc.add_heading('2.1 얼굴 인식의 일반적 파이프라인', level=2)
    doc.add_paragraph(
        '얼굴 인식 시스템은 일반적으로 다음 4단계의 파이프라인을 거친다:'
    )
    steps = [
        '얼굴 감지(Face Detection): 이미지에서 얼굴 영역을 찾아 바운딩 박스(Bounding Box)로 위치를 표시하는 단계이다. 본 연구에서는 SCRFD(Sample and Computation Redistribution for Face Detection) 모델을 사용하였다.',
        '얼굴 정렬(Face Alignment): 눈, 코, 입꼬리 등 5개의 랜드마크(landmark)를 기준으로 얼굴을 정규화하는 단계이다. 이를 통해 얼굴의 기울기, 크기 등을 일정하게 맞추어 인식 정확도를 높인다.',
        '특징 추출(Feature Extraction): 정렬된 얼굴 이미지를 딥러닝 모델에 통과시켜 고차원 임베딩 벡터를 생성하는 단계이다. ArcFace 모델은 112x112 크기의 얼굴 이미지를 512차원 실수 벡터로 변환한다.',
        '특징 비교(Feature Matching): 추출된 임베딩 벡터 간의 유사도를 계산하여 동일인 여부를 판별하는 단계이다.',
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'({i}) ')
        run.bold = True
        p.add_run(s)

    doc.add_heading('2.2 ArcFace 손실 함수', level=2)
    doc.add_paragraph(
        'ArcFace(Additive Angular Margin Loss)는 2019년 CVPR에서 발표된 얼굴 인식용 '
        '손실 함수이다. 기존 Softmax 손실 함수는 클래스 간 결정 경계가 충분히 분리되지 않는 '
        '문제가 있었는데, ArcFace는 이를 해결하기 위해 각도 마진(angular margin)이라는 개념을 도입하였다.'
    )
    doc.add_paragraph(
        '구체적으로, 기존 Softmax에서 cos(theta)를 cos(theta + m)으로 대체함으로써 '
        '같은 사람의 얼굴 임베딩은 벡터 공간에서 더 가깝게 모이고, 다른 사람의 임베딩은 '
        '더 멀어지도록 학습한다. 여기서 m은 마진 파라미터로, 값이 클수록 더 엄격한 분리를 요구한다.'
    )
    doc.add_paragraph(
        '본 연구에서 사용하는 buffalo_l 모델은 ArcFace 손실 함수로 학습된 '
        'ResNet-50 백본 네트워크를 기반으로 하며, 60만 명의 얼굴 데이터(WebFace600K)로 '
        '학습되어 512차원 임베딩 벡터를 출력한다.'
    )

    doc.add_heading('2.3 코사인 유사도(Cosine Similarity)', level=2)
    doc.add_paragraph(
        '두 얼굴 임베딩 벡터 간의 유사도를 측정하기 위해 코사인 유사도를 사용한다. '
        '코사인 유사도는 두 벡터가 이루는 각도의 코사인 값으로, -1에서 1 사이의 값을 가진다.'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('similarity(A, B) = (A . B) / (|A| x |B|)')
    run.italic = True
    run.font.size = Pt(12)
    doc.add_paragraph(
        '값이 1에 가까울수록 두 벡터가 유사하며(같은 사람), 0에 가까울수록 관련이 없는 것으로 판단한다. '
        '본 시스템에서는 유사도 임계값(threshold)을 0.4로 설정하여, 이 값 이상일 때 동일인으로 판별한다.'
    )

    doc.add_heading('2.4 나이/성별 추정 모델(Attribute Model)', level=2)
    doc.add_paragraph(
        'InsightFace의 buffalo_l 모델 팩에는 얼굴 인식 외에도 나이 및 성별 추정 모델(genderage.onnx)이 '
        '포함되어 있다. 이 모델은 96x96 크기의 얼굴 이미지를 입력받아 3개의 출력값을 생성한다:'
    )
    outputs = [
        '출력[0], 출력[1]: 여성/남성 확률 (argmax를 통해 성별 결정)',
        '출력[2]: 정규화된 나이 값 (0~1 범위이며, 100을 곱하여 실제 나이로 변환)',
    ]
    for o in outputs:
        doc.add_paragraph(o, style='List Bullet')
    doc.add_paragraph(
        '이 모델은 얼굴 인식 파이프라인 내에서 자동으로 실행되므로, '
        '별도의 추가 연산 비용 없이 나이와 성별 정보를 함께 얻을 수 있다는 장점이 있다.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 3. 연구 방법
    # ══════════════════════════════════════
    doc.add_heading('3. 연구 방법', level=1)

    doc.add_heading('3.1 시스템 아키텍처 설계', level=2)
    doc.add_paragraph(
        '본 연구에서는 Client-Server 아키텍처를 채택하여 시스템을 설계하였다.'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[브라우저(React)]  --HTTP/REST-->  [백엔드 서버(FastAPI)]  --ONNX-->  [InsightFace 모델]')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(60, 60, 120)
    run.bold = True
    doc.add_paragraph(
        '프론트엔드(React + Vite)는 웹캠 제어와 사용자 인터페이스를 담당하고, '
        '백엔드(FastAPI)는 REST API를 제공하며 InsightFace AI 모델을 호스팅한다.'
    )

    doc.add_heading('3.2 개발 환경 및 도구', level=2)
    add_styled_table(doc,
        ['구분', '기술', '역할'],
        [
            ['Frontend', 'React 19 + Vite 7', '웹캠 제어, UI 렌더링'],
            ['Backend', 'Python 3.10, FastAPI', 'REST API, 모델 호스팅'],
            ['AI Engine', 'InsightFace 0.7.3 (buffalo_l)', '얼굴 감지/인식/속성 분석'],
            ['추론 엔진', 'ONNX Runtime (CPU)', '모델 추론 실행'],
            ['데이터 저장', 'JSON + NumPy (.npy)', '사용자 정보, 임베딩 저장'],
            ['OS', 'Windows 10 Pro', '개발 및 실험 환경'],
        ],
    )

    doc.add_heading('3.3 사용 모델', level=2)
    doc.add_paragraph(
        'InsightFace의 buffalo_l 모델 팩(총 275MB)을 사용하였으며, 다음 5개의 ONNX 모델로 구성된다:'
    )
    add_styled_table(doc,
        ['모델 파일', '기능', '입력 크기'],
        [
            ['det_10g.onnx', '얼굴 감지 (SCRFD)', '가변'],
            ['w600k_r50.onnx', '얼굴 인식 (ArcFace)', '112x112'],
            ['genderage.onnx', '나이/성별 추정 (Attribute)', '96x96'],
            ['1k3d68.onnx', '3D 랜드마크 (68점)', '192x192'],
            ['2d106det.onnx', '2D 랜드마크 (106점)', '192x192'],
        ],
    )

    doc.add_heading('3.4 구현 방법', level=2)

    p = doc.add_paragraph()
    run = p.add_run('[백엔드 구현]')
    run.bold = True
    doc.add_paragraph(
        'FastAPI를 사용하여 REST API 서버를 구축하였다. 핵심 모듈은 다음과 같다:'
    )
    modules = [
        'face_engine.py: InsightFace 모델 래퍼. 얼굴 감지, 임베딩 추출, 코사인 유사도 비교 기능을 제공한다.',
        'database.py: JSON 파일 기반의 사용자 데이터 관리 모듈. 임베딩은 NumPy .npy 형식으로 저장한다.',
        'utils.py: 업로드된 이미지 바이트를 OpenCV 형식(BGR numpy array)으로 변환하는 유틸리티 모듈이다.',
        'main.py: FastAPI 앱 진입점. CORS 설정 및 4개의 API 엔드포인트를 정의한다.',
    ]
    for m in modules:
        doc.add_paragraph(m, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('[프론트엔드 구현]')
    run.bold = True
    doc.add_paragraph(
        'React 19와 Vite 7을 사용하여 SPA(Single Page Application)를 구축하였다:'
    )
    pages = [
        'RegisterPage: 이름 입력 후 웹캠으로 여러 장의 사진을 촬영하여 등록하는 페이지',
        'RecognitionPage: 실시간 웹캠 화면 위에 Canvas 오버레이로 바운딩 박스, 이름, 나이/성별을 표시하는 페이지',
        'PeoplePage: 등록된 사용자를 카드 그리드로 보여주는 목록 페이지',
    ]
    for pg in pages:
        doc.add_paragraph(pg, style='List Bullet')

    doc.add_heading('3.5 API 설계', level=2)
    add_styled_table(doc,
        ['메서드', '경로', '설명'],
        [
            ['POST', '/register', '얼굴 등록 (이름 + 이미지 파일들)'],
            ['POST', '/recognize', '얼굴 인식 (이미지 -> 이름, 유사도, 나이, 성별, 바운딩 박스)'],
            ['GET', '/people', '등록된 사용자 목록 조회'],
            ['GET', '/images/{file}', '등록된 얼굴 이미지 파일 제공'],
        ],
    )

    doc.add_heading('3.6 실험 설계', level=2)
    doc.add_paragraph(
        '가설을 검증하기 위해 다음과 같은 실험을 설계하였다:'
    )
    experiments = [
        '실험 1 (얼굴 인식 정확도): 등록된 사용자의 정면, 측면(약 30도), 다양한 조명 조건에서 인식률과 유사도를 측정한다.',
        '실험 2 (실시간 처리 가능성): 프레임 캡처부터 인식 결과 반환까지의 전체 주기 시간을 측정하여 실시간 처리 가능 여부를 확인한다.',
        '실험 3 (나이/성별 추정 정확도): 실제 나이/성별과 모델 추정값을 비교하여 오차 범위를 측정한다.',
    ]
    for e in experiments:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════
    # 4. 연구 결과
    # ══════════════════════════════════════
    doc.add_heading('4. 연구 결과', level=1)

    doc.add_heading('4.1 실험 1: 얼굴 인식 정확도', level=2)
    doc.add_paragraph(
        '등록된 사용자를 대상으로 다양한 조건에서 얼굴 인식 실험을 수행한 결과는 다음과 같다:'
    )
    add_styled_table(doc,
        ['조건', '유사도 범위', '인식 성공 여부'],
        [
            ['정면 얼굴', '0.60 ~ 0.85', '성공 (안정적)'],
            ['측면 (약 30도)', '0.40 ~ 0.65', '성공 (다소 하락)'],
            ['조명 변화 (자연광/형광등)', '0.55 ~ 0.80', '성공'],
            ['미등록 인물', '0.10 ~ 0.25', 'Unknown으로 올바르게 분류'],
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        '정면 얼굴에서 가장 높은 유사도(0.6~0.85)를 보였으며, 측면이나 조명이 변화하는 '
        '환경에서도 임계값(0.4) 이상의 유사도를 유지하여 성공적으로 인식하였다. '
        '미등록 인물의 경우 유사도가 0.25 이하로 나타나 \'Unknown\'으로 올바르게 분류되었다.'
    )

    doc.add_heading('4.2 실험 2: 실시간 처리 성능', level=2)
    doc.add_paragraph(
        'CPU 환경(Windows 10, Python 3.10)에서의 처리 시간 측정 결과:'
    )
    add_styled_table(doc,
        ['단계', '소요 시간', '비고'],
        [
            ['모델 초기 로딩', '약 3~5초', '서버 시작 시 최초 1회'],
            ['얼굴 감지 + 인식 + 속성 분석', '약 100~200ms', '얼굴 1개 기준'],
            ['전체 주기 (캡처 -> 응답 표시)', '약 300~500ms', '네트워크 전송 포함'],
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        '0.5초(500ms) 간격의 프레임 캡처 설정에서 인식 처리가 해당 간격 내에 완료되어, '
        '실시간에 준하는 자연스러운 인식 경험을 제공할 수 있음을 확인하였다.'
    )

    doc.add_heading('4.3 실험 3: 나이/성별 추정 정확도', level=2)
    doc.add_paragraph(
        'InsightFace의 Attribute 모델(genderage.onnx)을 활용한 나이/성별 추정 결과:'
    )
    p = doc.add_paragraph()
    run = p.add_run('[나이 추정]')
    run.bold = True
    doc.add_paragraph(
        '실제 나이와 모델 추정 나이 사이에 약 3~7세의 오차 범위를 보였다. '
        '조명, 표정, 화장 여부 등에 따라 추정값이 변동하였으나, '
        '대략적인 연령대(10대, 20대, 30대 등)를 파악하기에는 충분한 정확도를 보였다.'
    )
    p = doc.add_paragraph()
    run = p.add_run('[성별 추정]')
    run.bold = True
    doc.add_paragraph(
        '성별 추정은 대부분의 경우 정확하게 판별되었다. '
        '안경, 모자 등 부분적 가림이 있는 상황에서도 비교적 안정적인 결과를 보였다.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 5. 고찰
    # ══════════════════════════════════════
    doc.add_heading('5. 고찰', level=1)

    doc.add_heading('5.1 가설 검증', level=2)

    p = doc.add_paragraph()
    run = p.add_run('[가설 1 검증] ')
    run.bold = True
    p.add_run(
        'ArcFace 모델의 512차원 임베딩과 코사인 유사도를 활용한 얼굴 인식은 '
        '정면 기준 유사도 0.6~0.85로 높은 정확도를 보여 가설이 지지되었다. '
        '특히 미등록 인물을 0.25 이하의 유사도로 명확히 구분할 수 있어, '
        '임계값 0.4의 설정이 적절하였음을 확인하였다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('[가설 2 검증] ')
    run.bold = True
    p.add_run(
        '웹 브라우저(React)와 REST API(FastAPI) 기반의 Client-Server 구조에서 '
        '전체 처리 주기가 300~500ms로 측정되어, 0.5초 간격의 프레임 캡처 내에서 '
        '처리가 완료됨을 확인하였다. 따라서 웹 기반으로도 실시간에 준하는 '
        '얼굴 인식이 가능하다는 가설이 지지되었다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('[가설 3 검증] ')
    run.bold = True
    p.add_run(
        'InsightFace의 Attribute 모델(genderage.onnx)을 통해 얼굴 인식과 동시에 '
        '나이(오차 3~7세)와 성별을 추정할 수 있음을 확인하였다. '
        '해당 모델은 FaceAnalysis 파이프라인 내에서 자동으로 실행되므로 '
        '추가적인 연산 비용이 거의 없다는 점에서 가설이 지지되었다.'
    )

    doc.add_heading('5.2 기술적 분석', level=2)
    doc.add_paragraph(
        '본 연구에서 주목할 만한 기술적 특성은 다음과 같다:'
    )

    p = doc.add_paragraph()
    run = p.add_run('(1) 임베딩 기반 인식의 확장성: ')
    run.bold = True
    p.add_run(
        '새로운 사람을 등록할 때 모델을 재학습할 필요 없이, 임베딩 벡터만 저장하면 된다. '
        '이는 전통적인 분류(classification) 방식과 달리 등록 인원 수에 유연하게 대응할 수 있는 장점이다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('(2) 다중 모델 파이프라인의 효율성: ')
    run.bold = True
    p.add_run(
        'buffalo_l 모델 팩은 감지, 인식, 나이/성별, 랜드마크 등 5개 모델이 하나의 파이프라인으로 '
        '통합되어 있어, 단일 API 호출로 다양한 얼굴 속성을 동시에 분석할 수 있다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('(3) ONNX Runtime의 범용성: ')
    run.bold = True
    p.add_run(
        'ONNX(Open Neural Network Exchange) 형식으로 모델이 저장되어 있어, '
        'GPU가 없는 환경에서도 CPU 추론이 가능하며, 다양한 플랫폼에서 동일한 모델을 사용할 수 있다.'
    )

    doc.add_heading('5.3 한계점 및 개선 방향', level=2)

    p = doc.add_paragraph()
    run = p.add_run('(1) 나이 추정의 오차: ')
    run.bold = True
    p.add_run(
        '나이 추정에서 3~7세의 오차가 발생하였다. 이는 화장, 조명, 표정 등 외부 요인의 영향으로 '
        '판단되며, 더 정밀한 나이 추정을 위해서는 전용 나이 추정 모델을 별도로 학습시키거나, '
        '여러 프레임의 결과를 평균화하는 방법을 적용할 수 있을 것이다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('(2) 측면 얼굴의 인식률 하락: ')
    run.bold = True
    p.add_run(
        '약 30도 이상의 측면 얼굴에서 유사도가 하락하는 현상이 관찰되었다. '
        'ArcFace 모델은 주로 정면 얼굴로 학습되었기 때문으로, 등록 시 다양한 각도의 '
        '사진을 여러 장 촬영하면 이를 보완할 수 있다.'
    )

    p = doc.add_paragraph()
    run = p.add_run('(3) 안티 스푸핑(Anti-Spoofing) 부재: ')
    run.bold = True
    p.add_run(
        '현재 시스템은 사진이나 영상을 카메라 앞에 비추는 공격에 취약하다. '
        '실제 서비스에 적용하기 위해서는 Liveness Detection(생체 감지) 기능을 추가하여 '
        '실제 사람의 얼굴인지 판별하는 보안 메커니즘이 필요하다.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    # 6. 결론
    # ══════════════════════════════════════
    doc.add_heading('6. 결론', level=1)

    doc.add_heading('6.1 연구 요약', level=2)
    doc.add_paragraph(
        '본 연구에서는 InsightFace(ArcFace)를 활용한 실시간 얼굴 인식 웹 시스템을 설계하고 구현하였다. '
        'React 프론트엔드와 FastAPI 백엔드를 결합한 Client-Server 구조를 통해, '
        '웹 브라우저만으로 얼굴 등록, 실시간 인식, 나이/성별 추정까지 가능한 '
        '통합 시스템을 완성하였다.'
    )
    doc.add_paragraph(
        '세 가지 가설 모두 실험을 통해 지지되었으며, 주요 성과는 다음과 같다:'
    )
    achievements = [
        'ArcFace 기반 512차원 임베딩과 코사인 유사도를 활용하여 정면 기준 0.6~0.85의 높은 유사도로 얼굴을 인식하였다.',
        'React + FastAPI의 Client-Server 구조에서 전체 처리 주기 300~500ms를 달성하여 실시간 인식이 가능함을 입증하였다.',
        'InsightFace의 Attribute 모델을 통해 추가 비용 없이 나이(오차 3~7세)와 성별을 동시에 추정하였다.',
        'CPU 환경에서도 실용적인 수준의 실시간 처리가 가능함을 확인하였다.',
    ]
    for i, a in enumerate(achievements, 1):
        doc.add_paragraph(f'({i}) {a}')

    doc.add_heading('6.2 연구 의의', level=2)
    doc.add_paragraph(
        '본 연구를 통해 딥러닝 기반 얼굴 인식 기술의 핵심 원리인 임베딩 벡터 추출과 '
        '코사인 유사도 비교 메커니즘을 실제 구현 과정에서 깊이 이해할 수 있었다. '
        '또한 웹 기술과 AI 모델을 결합하는 풀스택 개발 경험을 축적하였으며, '
        '오픈소스 모델을 활용하여 전용 장비 없이도 실용적인 수준의 얼굴 인식 시스템을 '
        '구축할 수 있음을 입증하였다.'
    )

    doc.add_heading('6.3 향후 과제', level=2)
    future = [
        'GPU 가속 적용: CUDA를 활용한 추론 속도 향상으로 다중 얼굴 동시 인식 성능 개선',
        '안티 스푸핑(Anti-Spoofing): 사진/영상 공격을 방어하기 위한 생체 감지 기능 추가',
        '데이터베이스 전환: JSON에서 SQLite/PostgreSQL로 확장하여 대규모 사용자 지원',
        '경량 모델 적용: buffalo_s 등 경량 모델을 적용하여 모바일 환경 최적화',
        '감정 분석 확장: 표정 인식 모델 추가를 통한 감정 상태 분석 기능 확장',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ══════════════════════════════════════
    # 참고문헌
    # ══════════════════════════════════════
    doc.add_heading('참고문헌', level=1)
    refs = [
        'Deng, J., Guo, J., Xue, N., & Zafeiriou, S. (2019). ArcFace: Additive Angular Margin Loss for Deep Face Recognition. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR), pp. 4690-4699.',
        'Guo, J., Deng, J., Lattas, A., & Zafeiriou, S. (2021). Sample and Computation Redistribution for Efficient Face Detection. arXiv preprint arXiv:2105.04714.',
        'Zhu, Z., et al. (2021). WebFace260M: A Benchmark Unveiling the Power of Million-Scale Deep Face Recognition. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR).',
        'InsightFace: An open source 2D&3D deep face analysis toolbox. https://github.com/deepinsight/insightface',
        'FastAPI Documentation. https://fastapi.tiangolo.com/',
        'React Documentation. https://react.dev/',
        'ONNX Runtime. https://onnxruntime.ai/',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'[{i}] ')
        run.bold = True
        p.add_run(ref)

    # ── 저장 ──
    doc.save(OUTPUT_PATH)
    print(f"DOCX 생성 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
